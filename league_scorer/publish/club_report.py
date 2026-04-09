"""Generate club-level reports for a season.

Implements Club Reports v1.0.0 — single DOCX containing a section per
club with race-by-race team breakdowns, category summaries, unique
runner counts and top-5 runners. The DOCX formatting re-uses the
existing helpers in `report_writer.py` so the output matches other
reports produced by the application.
"""
from __future__ import annotations

from datetime import datetime, timezone
from pathlib import Path
import json
import logging

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt

from league_scorer.main import LeagueScorer
from league_scorer.season_aggregation import build_individual_season, build_team_season
from scripts.run_full_autopilot import _resolve_data_root, _season_paths
from league_scorer.output_layout import build_output_paths
from league_scorer.report_writer import (
    _apply_document_font_defaults,
    _section_heading,
    _sub_heading,
    _cell_text,
    _build_footer,
    _set_cell_bg,
    _NAVY_HEX,
    _WHITE_RGB,
)

log = logging.getLogger(__name__)


def _competition_ranks_for_race(team_results: list, division: int) -> dict:
    div = [t for t in team_results if t.division == division]
    with_runners = [t for t in div if t.team_score > 0]
    with_runners.sort(key=lambda t: t.team_score, reverse=True)
    pos_map: dict = {}
    rank = 1
    i = 0
    n = len(with_runners)
    while i < n:
        j = i + 1
        while j < n and with_runners[j].team_score == with_runners[i].team_score:
            j += 1
        for k in range(i, j):
            key = (with_runners[k].preferred_club, with_runners[k].team_id)
            pos_map[key] = rank
        rank += j - i
        i = j
    return pos_map


def _category_match(cat: str, token: str) -> bool:
    return token.lower() in (cat or "").lower()


def _best_runner(recs: list, club: str, gender: str, category_token: str):
    matches = [r for r in recs if r.preferred_club == club and r.gender == gender and _category_match(r.category, category_token)]
    if not matches:
        return None
    matches.sort(key=lambda r: (r.total_points, r.races_completed), reverse=True)
    return matches[0]


def _team_average_breakdown(a_rec, b_rec) -> tuple[str, str]:
    a_male_scores = [t.men_score for t in a_rec.race_results.values()] if a_rec else []
    a_female_scores = [t.women_score for t in a_rec.race_results.values()] if a_rec else []
    a_male_avg = sum(a_male_scores) / len(a_male_scores) if a_male_scores else 0.0
    a_female_avg = sum(a_female_scores) / len(a_female_scores) if a_female_scores else 0.0
    b_male_scores = [t.men_score for t in b_rec.race_results.values()] if b_rec else []
    b_female_scores = [t.women_score for t in b_rec.race_results.values()] if b_rec else []
    b_male_avg = sum(b_male_scores) / len(b_male_scores) if b_male_scores else 0.0
    b_female_avg = sum(b_female_scores) / len(b_female_scores) if b_female_scores else 0.0
    a_text = f"Team A average scores: Male {a_male_avg:.1f}, Female {a_female_avg:.1f}"
    b_text = f"Team B average scores: Male {b_male_avg:.1f}, Female {b_female_avg:.1f}"
    return a_text, b_text


def _build_club_header(doc: Document, year: int, club: str, images_dir: Path | None) -> None:
    tbl = doc.add_table(rows=1, cols=2)
    logo_cell = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]
    logo_cell.width = Cm(3.0)
    title_cell.width = Cm(15.0)

    _set_cell_bg(logo_cell, _NAVY_HEX)
    _set_cell_bg(title_cell, _NAVY_HEX)

    shield_path = (images_dir / "WRRL shield concept.png") if images_dir else None
    if shield_path and shield_path.exists():
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(shield_path), width=Cm(2.8))

    title_cell.paragraphs[0].clear()
    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.left_indent = Cm(0.4)
    p1.paragraph_format.space_before = Pt(8)
    r1 = p1.add_run(f"WRRL | Season Summary {year}")
    r1.bold = True
    r1.font.size = Pt(20)
    r1.font.color.rgb = _WHITE_RGB

    p2 = title_cell.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    r2 = p2.add_run(club.upper())
    r2.bold = True
    r2.font.size = Pt(18)
    r2.font.color.rgb = _WHITE_RGB

    # No lower subtitle line; header is intentionally minimal.


def _best_single_race_performers(recs: list, top_n: int = 2) -> list[tuple[str, int, int]]:
    scored = []
    for r in recs:
        for race_num, pts in r.race_points.items():
            scored.append((r.name, race_num, pts))
    scored.sort(key=lambda item: item[2], reverse=True)
    return scored[:top_n]


def _most_consistent_runners(recs: list, min_races: int = 3, top_n: int = 2) -> list[tuple[str, float, int]]:
    candidates = [r for r in recs if r.races_completed >= min_races]
    candidates.sort(key=lambda r: (r.total_points / r.races_completed if r.races_completed else 0, r.races_completed), reverse=True)
    return [(r.name, r.total_points / r.races_completed if r.races_completed else 0, r.races_completed) for r in candidates[:top_n]]


def _most_races_runners(recs: list, top_n: int = 2) -> list[tuple[str, int, int]]:
    sorted_recs = sorted(recs, key=lambda r: (r.races_completed, r.total_points), reverse=True)
    return [(r.name, r.races_completed, r.total_points) for r in sorted_recs[:top_n]]


def _new_returning_proxy(recs: list) -> tuple[int, int]:
    one_race = sum(1 for r in recs if r.races_completed == 1)
    returning = sum(1 for r in recs if r.races_completed >= 3)
    return one_race, returning


def _shade_cell(cell, color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)


def _make_pie_chart(path: Path, data: dict[str, int], title: str, size: int = 240) -> None:
    labels = [k for k, v in data.items() if v > 0]
    values = [v for v in data.values() if v > 0]
    if not values:
        labels = ["None"]
        values = [1]

    colors = [
        (51, 102, 204),
        (0, 153, 68),
        (255, 140, 0),
        (204, 0, 51),
        (128, 0, 128),
        (255, 102, 178),
        (0, 153, 204),
    ]
    if title.startswith("Gender"):
        colors = [(0, 102, 204), (204, 0, 102)]
    total = sum(values)
    angles = [360 * v / total for v in values]

    legend_width = 180
    img_width = size + legend_width
    img_height = size + 50
    img = Image.new("RGB", (img_width, img_height), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("arial.ttf", 14)
    except Exception:
        font = ImageFont.load_default()

    draw.text((20, 10), title, fill="black", font=font)
    bbox = [20, 40, size - 20, size - 20]
    start = 0
    for idx, angle in enumerate(angles):
        draw.pieslice(bbox, start, start + angle, fill=colors[idx % len(colors)])
        start += angle

    legend_x = size + 10
    legend_y = 40
    for idx, label in enumerate(labels):
        pct = int(round(data[label] / total * 100)) if total else 0
        text = f"{label}: {data[label]} ({pct}%)"
        draw.rectangle([legend_x, legend_y, legend_x + 14, legend_y + 14], fill=colors[idx % len(colors)])
        draw.text((legend_x + 18, legend_y), text, fill="black", font=font)
        legend_y += 22

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path, format="PNG")


def _sanitize_filename(value: str) -> str:
    return "".join(c if c.isalnum() else "_" for c in value).strip("_")[:80]


def _insert_charts(doc: Document, output_dir: Path, club: str, category_counts: dict[str, int], category_points: dict[str, int], gender_counts: dict[str, int], gender_points: dict[str, int]) -> None:
    club_key = _sanitize_filename(club)
    charts_dir = output_dir / "publish" / "docx" / "club-reports" / "charts"
    charts_dir.mkdir(parents=True, exist_ok=True)
    category_count_path = charts_dir / f"{club_key}_category_count.png"
    category_points_path = charts_dir / f"{club_key}_category_points.png"
    gender_count_path = charts_dir / f"{club_key}_gender_count.png"
    gender_points_path = charts_dir / f"{club_key}_gender_points.png"

    _make_pie_chart(category_count_path, category_counts, "Category runners")
    _make_pie_chart(category_points_path, category_points, "Category points")
    _make_pie_chart(gender_count_path, gender_counts, "Gender runners")
    _make_pie_chart(gender_points_path, gender_points, "Gender points")

    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"

    chart_labels = [
        ("Runners by Gender", gender_count_path),
        ("Runners by Category", category_count_path),
        ("Points by Gender", gender_points_path),
        ("Points by Category", category_points_path),
    ]

    for row_idx in range(2):
        for col_idx in range(2):
            label, path = chart_labels[row_idx * 2 + col_idx]
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            title_para = cell.add_paragraph()
            title_run = title_para.add_run(label)
            title_run.bold = True
            title_run.font.size = Pt(10)
            picture_para = cell.add_paragraph()
            picture_run = picture_para.add_run()
            picture_run.add_picture(str(path), width=Inches(2.0))


def generate_club_reports(year: int, data_root: Path | None, report_dir: Path) -> int:
    generated_at = datetime.now(timezone.utc).isoformat()

    data_root_resolved = _resolve_data_root(data_root)
    if data_root_resolved is None:
        print("No data root configured. Set Data Root before generating club reports.")
        return 1

    input_dir, output_dir = _season_paths(data_root_resolved, year)

    scorer = LeagueScorer(input_dir=input_dir, output_dir=output_dir, year=year)
    warnings = scorer.run()

    male_recs, female_recs = build_individual_season(scorer.all_race_runners)
    div1_teams, div2_teams = build_team_season(scorer.all_race_teams, scorer.club_info)

    highest = max(scorer.all_race_runners) if scorer.all_race_runners else 0

    doc = Document()
    _apply_document_font_defaults(doc)

    from docx.enum.section import WD_ORIENT
    from docx.shared import Cm

    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(1.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    images_dir = Path(__file__).parent.parent / "images"

    clubs = sorted(scorer.club_info.keys(), key=lambda s: s.lower())
    all_race_teams = scorer.all_race_teams

    for club in clubs:
        _build_club_header(doc, year, club, images_dir)

        _sub_heading(doc, "Summary")
        tbl = doc.add_table(rows=0, cols=2)
        tbl.style = "Table Grid"
        a_rec = next((t for t in div1_teams + div2_teams if t.preferred_club == club and t.team_id == "A"), None)
        b_rec = next((t for t in div1_teams + div2_teams if t.preferred_club == club and t.team_id == "B"), None)
        total_club_points = (a_rec.total_points if a_rec else 0) + (b_rec.total_points if b_rec else 0)
        row = tbl.add_row()
        _cell_text(row.cells[0], "Total club points")
        _cell_text(
            row.cells[1],
            f"{total_club_points} ({a_rec.total_points if a_rec else 0} Team A, {b_rec.total_points if b_rec else 0} Team B)",
        )
        row = tbl.add_row()
        _cell_text(row.cells[0], "A team overall position")
        _cell_text(row.cells[1], str(a_rec.position if a_rec and a_rec.position else ""))

        unique_names = set()
        for race_list in scorer.all_race_runners.values():
            for r in race_list:
                if r.preferred_club == club and r.eligible:
                    unique_names.add(r.name.lower())
        row = tbl.add_row()
        _cell_text(row.cells[0], "Unique runners (season)")
        _cell_text(row.cells[1], str(len(unique_names)))

        for idx, row in enumerate(tbl.rows, start=0):
            if idx % 2 == 1:
                for cell in row.cells:
                    _shade_cell(cell, "EFEFEF")

        combined_recs = [r for r in male_recs + female_recs if r.preferred_club == club]
        combined_recs.sort(key=lambda r: (r.total_points, r.races_completed), reverse=True)

        category_counts: dict[str, int] = {}
        category_points: dict[str, int] = {}
        for cat in ["Jun", "Sen", "V40", "V50", "V60", "V70"]:
            category_counts[cat] = sum(
                1 for r in combined_recs if _category_match(r.category, cat)
            )
            category_points[cat] = sum(
                r.total_points for r in combined_recs if _category_match(r.category, cat)
            )
        gender_counts = {
            "Male": sum(1 for r in combined_recs if r.gender == "M"),
            "Female": sum(1 for r in combined_recs if r.gender == "F"),
        }
        gender_points = {
            "Male": sum(r.total_points for r in combined_recs if r.gender == "M"),
            "Female": sum(r.total_points for r in combined_recs if r.gender == "F"),
        }

        _sub_heading(doc, "Race-by-Race Team Results (A / B)")
        race_tbl = doc.add_table(rows=1, cols=7)
        race_tbl.style = "Table Grid"
        hdr = race_tbl.rows[0].cells
        _cell_text(hdr[0], "Race #")
        _cell_text(hdr[1], "Team A — Score")
        _cell_text(hdr[2], "Team A — Pos")
        _cell_text(hdr[3], "Team A — Pts")
        _cell_text(hdr[4], "Team B — Score")
        _cell_text(hdr[5], "Team B — Pos")
        _cell_text(hdr[6], "Team B — Pts")

        for race_num in sorted(scorer.all_race_runners.keys()):
            teams = all_race_teams.get(race_num, [])
            pos_map_div1 = _competition_ranks_for_race(teams, division=(a_rec.division if a_rec else 1))
            pos_map_div2 = _competition_ranks_for_race(teams, division=(b_rec.division if b_rec else 2))

            a_team = next((t for t in teams if t.preferred_club == club and t.team_id == "A"), None)
            b_team = next((t for t in teams if t.preferred_club == club and t.team_id == "B"), None)

            row = race_tbl.add_row().cells
            _cell_text(row[0], str(race_num))
            _cell_text(row[1], str(a_team.team_score if a_team else ""))
            _cell_text(row[2], str(pos_map_div1.get((club, "A"), "")))
            _cell_text(row[3], str(a_team.team_points if a_team else ""))
            _cell_text(row[4], str(b_team.team_score if b_team else ""))
            _cell_text(row[5], str(pos_map_div2.get((club, "B"), "")))
            _cell_text(row[6], str(b_team.team_points if b_team else ""))

        # Shade the race-by-race header row and alternate rows
        for cell in race_tbl.rows[0].cells:
            _shade_cell(cell, "D9E1F2")
        for idx, row in enumerate(race_tbl.rows[1:], start=1):
            if idx % 2 == 0:
                for cell in row.cells:
                    _shade_cell(cell, "EFEFEF")

        _sub_heading(doc, "Top 5 Runners")
        combined_recs = [r for r in male_recs + female_recs if r.preferred_club == club]
        combined_recs.sort(key=lambda r: (r.total_points, r.races_completed), reverse=True)
        top5 = combined_recs[:5]
        t_tbl = doc.add_table(rows=1, cols=4)
        t_tbl.style = "Table Grid"
        hdr = t_tbl.rows[0].cells
        _cell_text(hdr[0], "Pos")
        _cell_text(hdr[1], "Name")
        _cell_text(hdr[2], "Total Points")
        _cell_text(hdr[3], "Races")
        for rec in top5:
            rrow = t_tbl.add_row().cells
            _cell_text(rrow[0], str(rec.position))
            _cell_text(rrow[1], rec.name)
            _cell_text(rrow[2], str(rec.total_points))
            _cell_text(rrow[3], str(rec.races_completed))

        for cell in t_tbl.rows[0].cells:
            _shade_cell(cell, "D9E1F2")
        for idx, row in enumerate(t_tbl.rows[1:], start=1):
            if idx % 2 == 0:
                for cell in row.cells:
                    _shade_cell(cell, "EFEFEF")

        _sub_heading(doc, "Top Category Performers")
        leader_tbl = doc.add_table(rows=1, cols=3)
        leader_tbl.style = "Table Grid"
        header = leader_tbl.rows[0].cells
        _cell_text(header[0], "Category")
        _cell_text(header[1], "Top Male")
        _cell_text(header[2], "Top Female")
        categories = ["Jun", "Sen", "V40", "V50", "V60", "V70"]
        for cat in categories:
            top_m = _best_runner(male_recs, club, "M", cat)
            top_f = _best_runner(female_recs, club, "F", cat)
            row = leader_tbl.add_row().cells
            _cell_text(row[0], cat)
            _cell_text(row[1], f"{top_m.name} ({top_m.total_points})" if top_m else "")
            _cell_text(row[2], f"{top_f.name} ({top_f.total_points})" if top_f else "")

        for cell in leader_tbl.rows[0].cells:
            _shade_cell(cell, "D9E1F2")
        for idx, row in enumerate(leader_tbl.rows[1:], start=1):
            if idx % 2 == 0:
                for cell in row.cells:
                    _shade_cell(cell, "EFEFEF")

        from league_scorer.report_writer import _page_break
        _page_break(doc)

        _sub_heading(doc, "Profile Charts")
        _insert_charts(doc, output_dir, club, category_counts, category_points, gender_counts, gender_points)

        one_race_count, returning_count = _new_returning_proxy(combined_recs)
        top_single = _best_single_race_performers(combined_recs, top_n=2)
        consistent = _most_consistent_runners(combined_recs, min_races=3, top_n=2)
        most_races = _most_races_runners(combined_recs, top_n=2)
        a_team_text, b_team_text = _team_average_breakdown(a_rec, b_rec)

        _sub_heading(doc, "Other Data")
        doc.add_paragraph(a_team_text, style="List Bullet")
        doc.add_paragraph(b_team_text, style="List Bullet")
        if most_races:
            most_races_text = "; ".join(f"{name} ({races} races, {pts} pts)" for name, races, pts in most_races)
            doc.add_paragraph(f"Runner(s) with most races: {most_races_text}", style="List Bullet")
        else:
            doc.add_paragraph("Runner(s) with most races: none", style="List Bullet")
        if top_single:
            top_text = "; ".join(f"{name} ({pts} pts, race {race_num})" for name, race_num, pts in top_single)
            doc.add_paragraph(f"Best single-race performer(s): {top_text}", style="List Bullet")
        else:
            doc.add_paragraph("Best single-race performer(s): none", style="List Bullet")
        if consistent:
            consistent_text = "; ".join(f"{name} ({avg:.1f} avg over {races} races)" for name, avg, races in consistent)
            doc.add_paragraph(f"Most consistent runners: {consistent_text}", style="List Bullet")
        else:
            doc.add_paragraph("Most consistent runners: not enough repeat race data", style="List Bullet")
        doc.add_paragraph(
            f"New/returning proxy: {one_race_count} one-race runner(s), {returning_count} runner(s) with 3+ races",
            style="List Bullet",
        )

        _page_break(doc)

    _build_footer(doc, year, include_page_numbers=True)

    output_paths = build_output_paths(output_dir)
    club_doc_dir = output_paths.publish_docx_league_updates_dir.parent / "club-reports"
    club_doc_dir.mkdir(parents=True, exist_ok=True)
    docx_path = club_doc_dir / f"club_reports_{year}.docx"
    doc.save(str(docx_path))
    log.info("Club report written: %s", docx_path)

    try:
        from docx2pdf import convert  # type: ignore
        pdf_path = docx_path.with_suffix(".pdf")
        try:
            convert(str(docx_path), str(pdf_path))
            log.info("Club report PDF written: %s", pdf_path)
        except Exception as exc:
            log.warning("Club report PDF conversion skipped — %s", exc)
    except Exception:
        log.info("docx2pdf not available; PDF conversion skipped for club reports")

    report_root = Path(report_dir) / f"year-{year}"
    report_root.mkdir(parents=True, exist_ok=True)
    payload = {
        "generated_at": generated_at,
        "year": year,
        "data_root": str(data_root_resolved),
        "docx": str(docx_path),
        "warnings": warnings,
    }
    json_path = report_root / "club_reports.json"
    md_path = report_root / "club_reports.md"
    json_path.write_text(json.dumps(payload, indent=2), encoding="utf-8")
    md_path.write_text("""# Club Reports

Generated: %s

See the DOCX output for full per-club details.
""" % (generated_at,), encoding="utf-8")

    return 0
