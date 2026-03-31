"""
report_writer.py — Branded DOCX + PDF season report.

Sections (combined league report)
----------------------------------
  1. Cover header      (WRRL shield + season title)
  2. League narrative  (season summary prose, auto end-of-season at Race 8)
  3. Division 1 Club Table
  4. Division 2 Club Table  (page 2)
  5. Top 20 Male Individual
  6. Top 20 Female Individual
  7. Category Leaders – top 3 per category (M & F side-by-side)
  8. Season overview stats
"""

import logging
import re
from collections import Counter
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

from .models import (
    RunnerRaceEntry,
    RunnerSeasonRecord,
    TeamRaceResult,
    TeamSeasonRecord,
    UnrecognisedClub,
)

log = logging.getLogger(__name__)

# ── Brand colours ──────────────────────────────────────────────────────────────
_NAVY_HEX    = "3a4658"
_GREEN_HEX   = "2d7a4a"
_WHITE_HEX   = "ffffff"
_ALT_HEX     = "eef2f7"   # alternating data row tint
_TOTAL_HEX   = "00b050"   # highlighted total-points column in division tables
_FONT_NAME   = "Calibri"
_SEASON_FINAL_RACE = 8

_NAVY_RGB   = RGBColor(0x3a, 0x46, 0x58)
_WHITE_RGB  = RGBColor(0xff, 0xff, 0xff)
_SUBHDR_RGB = RGBColor(0xa0, 0xb0, 0xc0)

# Award / status colours
_GOLD_HEX      = "c9a84c"   # 1st place — gold
_SILVER_HEX    = "a8b4c4"   # 2nd place — silver
_BRONZE_HEX    = "b87440"   # 3rd place — bronze
_PROMOTED_HEX  = "d0ead8"   # Div 2 promotion zone (subtle green tint)
_RELEGATED_HEX = "f5d8d8"   # Div 1 relegation zone (subtle red tint)

_CATEGORIES = ["Sen", "V40", "V50", "V60", "V70+"]

# Per-category badge background colours (hex without #)
_CAT_BADGE_HEX = {
    "Sen":  "2d7a4a",   # brand green
    "V40":  "1a5fa8",   # blue
    "V50":  "c05800",   # burnt orange
    "V60":  "5e3590",   # purple
    "V70+": "a06820",   # gold-brown
}

# row_style → (bg_hex, fg_RGBColor_or_None, bold, italic)
_ROW_STYLES = {
    "leader":   (_GREEN_HEX,    _WHITE_RGB, True,  False),
    "gold":     (_GOLD_HEX,     _NAVY_RGB,  True,  False),
    "silver":   (_SILVER_HEX,   _NAVY_RGB,  True,  False),
    "bronze":   (_BRONZE_HEX,   _WHITE_RGB, True,  False),
    "promoted": (_PROMOTED_HEX, _NAVY_RGB,  False, False),
    "relegated":(_RELEGATED_HEX,_NAVY_RGB,  False, True),
}


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    """Set (or replace) the background fill of a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for existing in tcPr.findall(qn("w:shd")):
        tcPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_para_bg(para, hex_color: str) -> None:
    """Shade an entire paragraph (fills to page margins — used for section bands)."""
    pPr = para._p.get_or_add_pPr()
    for existing in pPr.findall(qn("w:shd")):
        pPr.remove(existing)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    pPr.append(shd)


def _set_cell_borders(cell, **borders: str) -> None:
    """Set specific cell borders using Word border values."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tc_borders = tcPr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tcPr.append(tc_borders)

    for side, value in borders.items():
        existing = tc_borders.find(qn(f"w:{side}"))
        if existing is not None:
            tc_borders.remove(existing)
        border = OxmlElement(f"w:{side}")
        if value == "none":
            border.set(qn("w:val"), "none")
            border.set(qn("w:sz"), "0")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), "auto")
        else:
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), "6")
            border.set(qn("w:space"), "0")
            border.set(qn("w:color"), value)
        tc_borders.append(border)


def _cell_text(cell, text, bold=False, italic=False,
               color: RGBColor = None, size_pt=9, align="left") -> None:
    """Write text into a cell's first paragraph, clearing any prior content."""
    para = cell.paragraphs[0]
    para.clear()
    para.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    run = para.add_run("" if text is None or text == "" else str(text))
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = color


def _cell_name_with_club(cell, name: str, club: str,
                         bold=False, color: RGBColor = None,
                         align="left") -> None:
    """Write runner name with club underneath in a smaller muted line."""
    cell.text = ""

    p1 = cell.paragraphs[0]
    p1.alignment = {
        "left":   WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right":  WD_ALIGN_PARAGRAPH.RIGHT,
    }.get(align, WD_ALIGN_PARAGRAPH.LEFT)
    p1.paragraph_format.space_after = Pt(0)
    r1 = p1.add_run("" if not name else str(name))
    r1.bold = bold
    r1.font.size = Pt(9)
    if color:
        r1.font.color.rgb = color

    if club:
        p2 = cell.add_paragraph()
        p2.alignment = p1.alignment
        p2.paragraph_format.space_before = Pt(0)
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(str(club))
        r2.italic = True
        r2.font.size = Pt(8)
        r2.font.color.rgb = color if color == _WHITE_RGB else RGBColor(0x70, 0x80, 0x90)


def _append_page_number(paragraph) -> None:
    """Append a Word PAGE field to a paragraph."""
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


# ── Reusable table building blocks ─────────────────────────────────────────────

def _make_header_row(table, headers: list, widths_cm: list) -> None:
    """Style the pre-existing first row as a navy header band."""
    row = table.rows[0]
    for i, hdr in enumerate(headers):
        cell = row.cells[i]
        _set_cell_bg(cell, _NAVY_HEX)
        if i < len(widths_cm) and widths_cm[i]:
            cell.width = Cm(widths_cm[i])
        _cell_text(cell, hdr, bold=True, color=_WHITE_RGB, size_pt=8, align="center")


def _division_total_cell(cell, text, *, bold=False, italic=False,
                         size_pt=9, has_middle_line=True) -> None:
    """Render the division table total column with distinct fill and borders."""
    _set_cell_bg(cell, _TOTAL_HEX)
    _set_cell_borders(
        cell,
        left="none",
        top="none",
        bottom="FFFFFF" if has_middle_line else "none",
        right="none",
    )
    _cell_text(cell, text, bold=bold, italic=italic,
               color=_WHITE_RGB, size_pt=size_pt, align="center")


def _set_row_height(row, height_cm: float) -> None:
    """Apply a fixed row height so division tables keep a stable layout."""
    row.height = Cm(height_cm)
    row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _apply_document_font_defaults(doc: Document) -> None:
    """Set the generated document's styles to use a consistent font family."""
    for style in doc.styles:
        try:
            style.font.name = _FONT_NAME
        except Exception:
            continue


def _add_data_row(table, values: list, widths_cm: list,
                  alt=False, row_style: str = None, name_col=1) -> None:
    """Append a data row.  *row_style* keys into _ROW_STYLES."""
    row = table.add_row()
    if row_style and row_style in _ROW_STYLES:
        bg, fg, is_bold, is_italic = _ROW_STYLES[row_style]
    else:
        bg, fg, is_bold, is_italic = (_ALT_HEX if alt else _WHITE_HEX), None, False, False
    for i, val in enumerate(values):
        cell = row.cells[i]
        _set_cell_bg(cell, bg)
        if i < len(widths_cm) and widths_cm[i]:
            cell.width = Cm(widths_cm[i])
        align = "left" if i == name_col else "center"
        _cell_text(cell, val, bold=(is_bold or i == 0),
                   italic=is_italic, color=fg, size_pt=9, align=align)


# ── Section heading ────────────────────────────────────────────────────────────

def _page_break(doc: Document) -> None:
    """Insert a hard page break."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run().add_break(WD_BREAK.PAGE)


def _section_heading(doc: Document, text: str, space_before_pt=14) -> None:
    """Full-width green band with white bold caps text — matches the source image style."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.left_indent  = Pt(6)
    _set_para_bg(p, _GREEN_HEX)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _WHITE_RGB


def _sub_heading(doc: Document, text: str) -> None:
    """Smaller green-accented sub-section label."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = _NAVY_RGB


def _spacer(doc: Document, height_pt: float = 10) -> None:
    """Insert an empty paragraph for visual breathing room."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(height_pt)


def _set_table_cell_margins(tbl, top_pt=4, bottom_pt=4, left_pt=6, right_pt=6) -> None:
    """Set uniform top/bottom/left/right cell margins on a table."""
    tbl_elem = tbl._tbl
    tbl_pr = tbl_elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tbl_pr)
    for existing in tbl_pr.findall(qn("w:tblCellMar")):
        tbl_pr.remove(existing)
    mar = OxmlElement("w:tblCellMar")
    for side, val in (("top", top_pt), ("left", left_pt),
                      ("bottom", bottom_pt), ("right", right_pt)):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(int(val * 20)))   # twentieths of a point
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tbl_pr.append(mar)


def _remove_vertical_borders(tbl) -> None:
    """Remove only the vertical (insideV + left + right) borders; keep horizontal lines."""
    tbl_elem = tbl._tbl
    tbl_pr = tbl_elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tbl_pr)
    tbl_borders = tbl_pr.find(qn("w:tblBorders"))
    if tbl_borders is None:
        tbl_borders = OxmlElement("w:tblBorders")
        tbl_pr.append(tbl_borders)
    for name in ("left", "right", "insideV"):
        existing = tbl_borders.find(qn(f"w:{name}"))
        if existing is not None:
            tbl_borders.remove(existing)
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)


def _no_table_borders(tbl) -> None:
    """Remove all visible borders from a table."""
    tbl_elem = tbl._tbl
    tbl_pr = tbl_elem.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tbl_pr)
    for existing in tbl_pr.findall(qn("w:tblBorders")):
        tbl_pr.remove(existing)
    tbl_borders = OxmlElement("w:tblBorders")
    for name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{name}")
        el.set(qn("w:val"), "none")
        el.set(qn("w:sz"), "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        tbl_borders.append(el)
    tbl_pr.append(tbl_borders)


# ── Report sections ────────────────────────────────────────────────────────────

def _write_overview(doc: Document, highest: int, year: int,
                    male: List[RunnerSeasonRecord],
                    female: List[RunnerSeasonRecord],
                    unrec: List[UnrecognisedClub],
                    nonleague_filename: str = "") -> None:
    _section_heading(doc, "Season Overview", space_before_pt=6)

    def _write_summary_table(headers: list[str], widths_cm: list[float], rows: list[tuple],
                             name_col: int = 0) -> None:
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = "Table Grid"
        _remove_vertical_borders(table)
        _make_header_row(table, headers, widths_cm)
        _set_table_cell_margins(table, top_pt=5, bottom_pt=5)
        for idx, row_data in enumerate(rows):
            _add_data_row(table, list(row_data), widths_cm, alt=idx % 2 == 1, name_col=name_col)

    headers    = ["", "Male", "Female", "Total"]
    widths_cm  = [6.0, 3.0, 3.0, 3.0]
    clubs_m = {r.preferred_club for r in male}
    clubs_f = {r.preferred_club for r in female}
    clubs_all = clubs_m | clubs_f

    unrec_count = len({u.raw_club_name for u in unrec})
    nonleague_note = (
        f"{unrec_count} club(s) — see {nonleague_filename}" if unrec_count
        else "None"
    )

    rows = [
        ("Season",             str(year),          "",               ""),
        ("Races processed",    str(highest),        "",               ""),
        ("Runners scored",     str(len(male)),       str(len(female)), str(len(male) + len(female))),
        ("Clubs scored",       str(len(clubs_m)),    str(len(clubs_f)),str(len(clubs_all))),
        ("Non-league clubs",   nonleague_note,       "",               ""),
    ]
    _write_summary_table(headers, widths_cm, rows)

    male_by_category = Counter(r.category for r in male)
    female_by_category = Counter(r.category for r in female)
    category_rows = [
        (
            cat,
            str(male_by_category.get(cat, 0)),
            str(female_by_category.get(cat, 0)),
            str(male_by_category.get(cat, 0) + female_by_category.get(cat, 0)),
        )
        for cat in _CATEGORIES
        if male_by_category.get(cat, 0) or female_by_category.get(cat, 0)
    ]
    if category_rows:
        _sub_heading(doc, "Runner Categories")
        _write_summary_table(["Category", "Male", "Female", "Total"], [4.0, 2.5, 2.5, 2.5], category_rows)

    club_counter = Counter()
    for record in male + female:
        if record.preferred_club:
            club_counter[record.preferred_club] += 1

    club_rows = [
        (club, str(count))
        for club, count in sorted(club_counter.items(), key=lambda item: (-item[1], item[0]))
    ]
    if club_rows:
        _sub_heading(doc, "Unique Runners By Club")
        _write_summary_table(["Club", "Unique Runners"], [10.5, 3.0], club_rows)


def _team_narrative_name(team: TeamSeasonRecord) -> str:
    """Compact team label for prose summaries."""
    return team.display_name.replace(" -- ", " ")


def _runner_narrative_name(record: RunnerSeasonRecord) -> str:
    """Compact runner label for prose summaries."""
    return f"{record.name} ({record.preferred_club})"


def _write_league_narrative(doc: Document,
                            highest_race: int,
                            year: int,
                            male_records: List[RunnerSeasonRecord],
                            female_records: List[RunnerSeasonRecord],
                            div1_teams: List[TeamSeasonRecord],
                            div2_teams: List[TeamSeasonRecord]) -> None:
    """Write a short season-level narrative ahead of the division tables."""
    _sub_heading(doc, "League Summary")

    div1 = sorted(div1_teams, key=lambda team: team.position)
    div2 = sorted(div2_teams, key=lambda team: team.position)
    male = sorted(male_records, key=lambda record: record.position)
    female = sorted(female_records, key=lambda record: record.position)

    div1_leader = div1[0] if div1 else None
    div1_runner_up = div1[1] if len(div1) > 1 else None
    div1_bottom = div1[-2:] if len(div1) >= 2 else div1
    div2_leader = div2[0] if div2 else None
    div2_promotion = div2[:2]
    male_leader = male[0] if male else None
    female_leader = female[0] if female else None

    if highest_race == _SEASON_FINAL_RACE:
        parts = [
            f"Race {highest_race} brings the {year} WRRL league season to its conclusion.",
        ]
        if div1_leader and div1_runner_up:
            parts.append(
                f"Division 1 is won by {_team_narrative_name(div1_leader)} on {div1_leader.total_points} points, "
                f"with {_team_narrative_name(div1_runner_up)} finishing second on {div1_runner_up.total_points}."
            )
        elif div1_leader:
            parts.append(
                f"Division 1 is led by {_team_narrative_name(div1_leader)} on {div1_leader.total_points} points."
            )

        if len(div2_promotion) == 2:
            parts.append(
                f"Promotion from Division 2 goes to {_team_narrative_name(div2_promotion[0])} and "
                f"{_team_narrative_name(div2_promotion[1])}, while "
                f"{', '.join(_team_narrative_name(team) for team in div1_bottom)} finish in the relegation places."
            )
        elif div2_leader:
            parts.append(
                f"Division 2 is topped by {_team_narrative_name(div2_leader)} on {div2_leader.total_points} points."
            )

        if male_leader and female_leader:
            parts.append(
                f"The individual tables are headed by {_runner_narrative_name(male_leader)} with {male_leader.total_points} points "
                f"and {_runner_narrative_name(female_leader)} with {female_leader.total_points}."
            )
    else:
        parts = [
            f"After {highest_race} races, the {year} WRRL season has reached its middle phase and the league tables are beginning to settle.",
        ]
        if div1_leader and div1_runner_up:
            parts.append(
                f"{_team_narrative_name(div1_leader)} currently leads Division 1 on {div1_leader.total_points} points, "
                f"with {_team_narrative_name(div1_runner_up)} closest on {div1_runner_up.total_points}."
            )
        elif div1_leader:
            parts.append(
                f"{_team_narrative_name(div1_leader)} currently leads Division 1 on {div1_leader.total_points} points."
            )

        if len(div2_promotion) == 2:
            parts.append(
                f"In Division 2, {_team_narrative_name(div2_promotion[0])} sets the pace and the current promotion places are held by "
                f"{_team_narrative_name(div2_promotion[0])} and {_team_narrative_name(div2_promotion[1])}."
            )
        elif div2_leader:
            parts.append(
                f"In Division 2, {_team_narrative_name(div2_leader)} sets the pace on {div2_leader.total_points} points."
            )

        if male_leader and female_leader:
            parts.append(
                f"The individual standings are led by {_runner_narrative_name(male_leader)} and "
                f"{_runner_narrative_name(female_leader)}, leaving the remaining races to decide titles, promotion and relegation."
            )

    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.add_run(" ".join(parts))
    run.font.size = Pt(10)
    run.font.color.rgb = _NAVY_RGB


def _write_club_table(doc: Document, teams: List[TeamSeasonRecord],
                      title: str, race_count: int,
                      is_div1: bool = False) -> None:
    """Render one club division table.

    Div 1: bottom 2 teams are highlighted as relegated (subtle red, italic).
    Div 2: position 1 = winner (green); position 2 = promoted (subtle teal).
    In both divisions position 1 always gets the green winner highlight.
    """
    races      = list(range(1, race_count + 1))
    headers    = ["Pos", "Club"] + [f"R{n}" for n in races] + ["Total"]
    name_col   = 1
    widths_cm  = [1.0, 5.5] + [1.2] * len(races) + [2.1]
    title_text = title.replace(" — Club Table", "").upper()

    sorted_teams = sorted(teams, key=lambda x: x.position)
    if not is_div1:
        sorted_teams = [team for team in sorted_teams if team.total_points > 0]
    total        = len(sorted_teams)
    RELEGATE_N   = 2   # bottom N in Div 1
    PROMOTE_N    = 2   # top N in Div 2 (pos 1 stays as leader)

    t = doc.add_table(rows=2, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = False
    _no_table_borders(t)

    title_row = t.rows[0]
    _set_row_height(title_row, 0.5)
    merged_title = title_row.cells[0]
    for cell in title_row.cells[1:]:
        merged_title = merged_title.merge(cell)
    _set_cell_bg(merged_title, _NAVY_HEX)
    _cell_text(merged_title, title_text, bold=True,
               color=_WHITE_RGB, size_pt=12, align="center")

    header_row = t.rows[1]
    _set_row_height(header_row, 0.5)
    for i, hdr in enumerate(headers):
        cell = header_row.cells[i]
        if i < len(widths_cm) and widths_cm[i]:
            cell.width = Cm(widths_cm[i])
        if hdr == "Total":
            _division_total_cell(cell, hdr, bold=True, size_pt=8, has_middle_line=True)
        else:
            _set_cell_bg(cell, _NAVY_HEX)
            _cell_text(cell, hdr, bold=True, color=_WHITE_RGB, size_pt=8, align="center")

    for i, team in enumerate(sorted_teams):
        if team.position == 1:
            style = "leader"
        elif is_div1 and i >= total - RELEGATE_N:
            style = "relegated"
        elif not is_div1 and i < PROMOTE_N:   # pos 2 in Div 2
            style = "promoted"
        else:
            style = None

        if style and style in _ROW_STYLES:
            bg, fg, is_bold, is_italic = _ROW_STYLES[style]
        else:
            bg, fg, is_bold, is_italic = (
                (_ALT_HEX if i % 2 == 1 else _WHITE_HEX), None, False, False
            )

        row = t.add_row()
        _set_row_height(row, 0.5)
        vals = [team.position, team.display_name]
        for n in races:
            rr = team.race_results.get(n)
            vals.append(rr.team_points if rr else "")

        for j, val in enumerate(vals):
            cell = row.cells[j]
            _set_cell_bg(cell, bg)
            if j < len(widths_cm) and widths_cm[j]:
                cell.width = Cm(widths_cm[j])
            align = "left" if j == name_col else "center"
            _cell_text(cell, val, bold=(is_bold or j == 0),
                       italic=is_italic, color=fg, size_pt=9, align=align)

        total_cell = row.cells[len(headers) - 1]
        if len(widths_cm) >= len(headers):
            total_cell.width = Cm(widths_cm[-1])
        _division_total_cell(
            total_cell,
            team.total_points,
            bold=is_bold,
            italic=is_italic,
            has_middle_line=(i < total - 1),
        )

    # Legend note below the table
    legend_p = doc.add_paragraph()
    legend_p.paragraph_format.space_before = Pt(3)
    legend_p.paragraph_format.space_after  = Pt(4)
    if is_div1:
        _legend_chip(legend_p, _RELEGATED_HEX, "Relegated")
    else:
        _legend_chip(legend_p, _PROMOTED_HEX, "Promoted")
        r = legend_p.add_run("  (top 2 clubs promoted to Division 1)")
        r.font.size = Pt(8)
        r.font.color.rgb = RGBColor(0x60, 0x70, 0x80)


def _legend_chip(para, hex_color: str, label: str) -> None:
    """Inline coloured square + text label for table legends."""
    # Coloured square via background-shaded run
    run = para.add_run("  ")
    run.font.size = Pt(10)
    rPr = run._r.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    rPr.append(shd)
    run2 = para.add_run(f" {label}")
    run2.font.size = Pt(8)
    run2.font.color.rgb = RGBColor(0x60, 0x70, 0x80)


_MEDAL_STYLES = {1: "gold", 2: "silver", 3: "bronze"}


def _write_individual_table(doc: Document, records: List[RunnerSeasonRecord],
                             title: str, race_count: int, top_n: int = 20) -> None:
    """Top-N individual table.  Positions 1/2/3 get gold/silver/bronze highlights."""
    _section_heading(doc, title)
    races     = list(range(1, race_count + 1))
    headers   = ["Pos", "Name", "Cat", "Total"] + [f"R{n}" for n in races]
    name_col  = 1
    widths_cm = [0.8, 5.8, 1.2, 1.3] + [1.0] * len(races)

    top = sorted(records, key=lambda r: r.position)[:top_n]

    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    _remove_vertical_borders(t)
    _make_header_row(t, headers, widths_cm)
    _set_table_cell_margins(t, top_pt=5, bottom_pt=5)

    for i, rec in enumerate(top):
        style = _MEDAL_STYLES.get(rec.position)
        if style:
            bg, fg, is_bold, is_italic = _ROW_STYLES[style]
        else:
            bg, fg, is_bold, is_italic = (
                (_ALT_HEX if i % 2 == 1 else _WHITE_HEX), None, False, False
            )

        vals = [rec.position, None, rec.category, rec.total_points]
        for n in races:
            vals.append(rec.race_points.get(n, ""))

        row = t.add_row()
        for j, val in enumerate(vals):
            cell = row.cells[j]
            _set_cell_bg(cell, bg)
            if j < len(widths_cm) and widths_cm[j]:
                cell.width = Cm(widths_cm[j])

            if j == name_col:
                _cell_name_with_club(
                    cell,
                    rec.name,
                    rec.preferred_club,
                    bold=is_bold,
                    color=fg,
                    align="left",
                )
                continue

            align = "left" if j == name_col else "center"
            _cell_text(cell, val, bold=(is_bold or j == 0),
                       italic=is_italic, color=fg, size_pt=9, align=align)


def _write_category_leaders(doc: Document,
                              male: List[RunnerSeasonRecord],
                              female: List[RunnerSeasonRecord],
                              top_n: int = 3) -> None:
    """Category top-N tables. Position 1 is highlighted as category champion (green).
    A narrow spacer column separates the male and female halves.
    """
    _section_heading(doc, "Category Leaders")
    # Layout: 3 male cols | 1 spacer | 3 female cols  = 7 columns
    headers   = ["#", "Male Runner", "Pts", "", "#", "Female Runner", "Pts"]
    widths_cm = [0.6, 6.1, 1.0, 0.4, 0.6, 6.1, 1.0]
    name_cols = (1, 5)
    spacer_col = 3

    for cat in _CATEGORIES:
        m_top = sorted([r for r in male   if r.category == cat], key=lambda r: r.position)[:top_n]
        f_top = sorted([r for r in female if r.category == cat], key=lambda r: r.position)[:top_n]
        if not m_top and not f_top:
            continue

        _sub_heading(doc, cat)
        t = doc.add_table(rows=1, cols=7)
        t.style = "Table Grid"
        _remove_vertical_borders(t)
        _make_header_row(t, headers, widths_cm)
        _set_table_cell_margins(t, top_pt=5, bottom_pt=5)

        for i in range(max(len(m_top), len(f_top))):
            m = m_top[i] if i < len(m_top) else None
            f = f_top[i] if i < len(f_top) else None
            vals = [
                m.position if m else "",
                None,
                m.total_points   if m else "",
                "",                               # spacer
                f.position if f else "",
                None,
                f.total_points   if f else "",
            ]
            # Position 1 = category champion — green leader row
            style = "leader" if i == 0 else None
            row = t.add_row()
            if style:
                bg, fg, is_bold, _ = _ROW_STYLES[style]
            else:
                bg = _ALT_HEX if i % 2 == 1 else _WHITE_HEX
                fg = None
                is_bold = False
            for j, val in enumerate(vals):
                cell = row.cells[j]
                # Spacer column: always white, no text
                if j == spacer_col:
                    _set_cell_bg(cell, _WHITE_HEX)
                    cell.width = Cm(widths_cm[j])
                    continue
                _set_cell_bg(cell, bg)
                if j < len(widths_cm) and widths_cm[j]:
                    cell.width = Cm(widths_cm[j])
                if j == 1:
                    _cell_name_with_club(
                        cell,
                        m.name if m else "",
                        m.preferred_club if m else "",
                        bold=is_bold,
                        color=fg,
                        align="left",
                    )
                    continue
                if j == 5:
                    _cell_name_with_club(
                        cell,
                        f.name if f else "",
                        f.preferred_club if f else "",
                        bold=is_bold,
                        color=fg,
                        align="left",
                    )
                    continue
                align = "left" if j in name_cols else "center"
                _cell_text(cell, val, bold=(is_bold or j in (0, 4)),
                           color=fg, size_pt=9, align=align)


# ── Per-race report ───────────────────────────────────────────────────────────────────────────────────────────────────────────────────────

def _extract_race_name(filepath: Path) -> str:
    """Strip 'Race N -- ' prefix from the filename stem."""
    m = re.match(r'^Race\s+\d+\s*[-\u2013\u2014]+\s*(.+)$', Path(filepath).stem, re.IGNORECASE)
    return m.group(1).strip() if m else Path(filepath).stem


def _build_race_header(doc: Document, race_name: str, race_num: int,
                       total_races: int, images_dir: Optional[Path]) -> None:
    """Big race-name header matching the WRRL scoring table card style."""
    tbl = doc.add_table(rows=1, cols=2)
    logo_cell  = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]
    logo_cell.width  = Cm(2.5)
    title_cell.width = Cm(15.5)
    _set_cell_bg(logo_cell, _NAVY_HEX)
    _set_cell_bg(title_cell, _NAVY_HEX)

    shield_path = (images_dir / "WRRL shield concept.png") if images_dir else None
    if shield_path and shield_path.exists():
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(shield_path), width=Cm(2.2))

    title_cell.paragraphs[0].clear()
    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.left_indent  = Cm(0.4)
    p1.paragraph_format.space_before = Pt(6)
    r1 = p1.add_run(race_name.upper())
    r1.bold = True
    r1.font.size = Pt(22)
    r1.font.color.rgb = _WHITE_RGB

    p2 = title_cell.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    r2 = p2.add_run(f"RACE {race_num} OF {total_races}")
    r2.bold = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = _SUBHDR_RGB

    accent = doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(6)
    _set_para_bg(accent, _GREEN_HEX)


def _write_race_individual(doc: Document,
                           runners: List[RunnerRaceEntry],
                           top_n: int = 10) -> None:
    """Top-N eligible finishers by overall finish time (both genders together)."""
    _section_heading(doc, "Individual Results", space_before_pt=14)

    eligible = sorted(
        [r for r in runners if r.eligible and r.points > 0],
        key=lambda r: r.time_seconds,
    )[:top_n]

    if not eligible:
        doc.add_paragraph("  No eligible results.")
        return

    headers   = ["Pos", "Name", "Club", "Gender", "Cat", "Time", "Points"]
    widths_cm = [1.0, 5.0, 4.5, 1.5, 1.5, 2.5, 2.0]
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    _remove_vertical_borders(t)
    _set_table_cell_margins(t, top_pt=5, bottom_pt=5)
    _make_header_row(t, headers, widths_cm)

    for i, r in enumerate(eligible):
        vals  = [i + 1, r.name, r.preferred_club or "",
                 r.gender, r.normalised_category, r.time_str, r.points]
        style = _MEDAL_STYLES.get(i + 1)
        _add_data_row(t, vals, widths_cm, alt=i % 2 == 1,
                      row_style=style, name_col=1)

    _spacer(doc, height_pt=6)


def _write_race_teams(doc: Document, team_results: List[TeamRaceResult]) -> None:
    """Team results grouped by division — Division 1 left, Division 2 right."""
    _section_heading(doc, "Club Results", space_before_pt=14)

    def _sorted_div(div: int) -> List[TeamRaceResult]:
        return sorted(
            [t for t in team_results if t.division == div],
            key=lambda x: (-x.team_points, -x.team_score),
        )

    div1_teams = _sorted_div(1)
    div2_teams = [t for t in _sorted_div(2) if t.team_score > 0]

    # 13 columns: 6 (Div 1) + 1 (spacer) + 6 (Div 2)
    # Fits A4 portrait 18cm usable: (18.0 - 0.4) / 2 = 8.8cm per side
    col_widths = [0.7, 3.3, 1.2, 1.2, 1.3, 1.1]   # Pos | Club | Men | Women | Score | Pts
    div_w      = 0.4
    DIV_COL    = 6
    COLS       = 13

    t = doc.add_table(rows=2, cols=COLS)
    t.style = "Table Grid"
    _remove_vertical_borders(t)
    _set_table_cell_margins(t, top_pt=5, bottom_pt=5)

    def _pos_list(teams: List[TeamRaceResult]) -> List[int]:
        """Competition-ranked positions (tied teams share position)."""
        out: List[int] = []
        rank = 1
        i = 0
        n = len(teams)
        while i < n:
            j = i + 1
            while j < n and teams[j].team_points == teams[i].team_points:
                j += 1
            for _ in range(i, j):
                out.append(rank)
            rank = j + 1
            i = j
        return out

    pos1 = _pos_list(div1_teams)
    pos2 = _pos_list(div2_teams)

    # Row 0: merged section title headers
    r0 = t.rows[0]
    merged_l = r0.cells[0].merge(r0.cells[5])
    _set_cell_bg(merged_l, _NAVY_HEX)
    _cell_text(merged_l, "Division 1", bold=True, color=_WHITE_RGB, size_pt=10, align="center")
    _set_cell_bg(r0.cells[DIV_COL], _WHITE_HEX)
    r0.cells[DIV_COL].width = Cm(div_w)
    merged_r = r0.cells[7].merge(r0.cells[12])
    _set_cell_bg(merged_r, _NAVY_HEX)
    _cell_text(merged_r, "Division 2", bold=True, color=_WHITE_RGB, size_pt=10, align="center")

    # Row 1: column sub-headers
    r1 = t.rows[1]
    col_hdrs = ["Pos", "Club", "Men", "Women", "Score", "Pts"]
    for j, (hdr, w) in enumerate(zip(col_hdrs, col_widths)):
        c = r1.cells[j]
        _set_cell_bg(c, _NAVY_HEX)
        c.width = Cm(w)
        _cell_text(c, hdr, bold=True, color=_WHITE_RGB, size_pt=8, align="center")
    _set_cell_bg(r1.cells[DIV_COL], _WHITE_HEX)
    r1.cells[DIV_COL].width = Cm(div_w)
    for j, (hdr, w) in enumerate(zip(col_hdrs, col_widths)):
        c = r1.cells[7 + j]
        _set_cell_bg(c, _NAVY_HEX)
        c.width = Cm(w)
        _cell_text(c, hdr, bold=True, color=_WHITE_RGB, size_pt=8, align="center")

    # Data rows
    for i in range(max(len(div1_teams), len(div2_teams), 1)):
        row = t.add_row()
        for teams, positions, offset in (
            (div1_teams, pos1, 0),
            (div2_teams, pos2, 7),
        ):
            team = teams[i] if i < len(teams) else None
            pos  = positions[i] if i < len(positions) else None
            style = "leader" if (team and pos == 1) else None
            bg, fg, is_bold, _ = (
                _ROW_STYLES[style] if style
                else ((_ALT_HEX if i % 2 == 1 else _WHITE_HEX), None, False, False)
            )
            club_label = (
                f"{team.preferred_club} ({team.team_id})" if team else ""
            )
            vals = [
                pos,
                club_label,
                team.men_score   if team else "",
                team.women_score if team else "",
                team.team_score  if team else "",
                team.team_points if team else "",
            ]
            for j, (val, w) in enumerate(zip(vals, col_widths)):
                c = row.cells[offset + j]
                _set_cell_bg(c, bg)
                c.width = Cm(w)
                align = "left" if j == 1 else "center"
                _cell_text(c, val, bold=(is_bold or j == 0), color=fg,
                           size_pt=9, align=align)

        _set_cell_bg(row.cells[DIV_COL], _WHITE_HEX)
        row.cells[DIV_COL].width = Cm(div_w)

    _spacer(doc, height_pt=6)


def _write_race_category_strip(doc: Document,
                               runners: List[RunnerRaceEntry]) -> None:
    """Category leaders — coloured badge strip, male and female champion per category."""
    _section_heading(doc, "Category Leaders", space_before_pt=14)

    # Collect best M + F per category
    m_champs: dict = {}
    f_champs: dict = {}
    for cat in _CATEGORIES:
        m_champs[cat] = next(iter(sorted(
            [r for r in runners if r.eligible and r.normalised_category == cat
             and r.gender == "M" and r.points > 0],
            key=lambda r: (-r.points, r.time_seconds),
        )), None)
        f_champs[cat] = next(iter(sorted(
            [r for r in runners if r.eligible and r.normalised_category == cat
             and r.gender == "F" and r.points > 0],
            key=lambda r: (-r.points, r.time_seconds),
        )), None)

    n_cats    = len(_CATEGORIES)
    label_w   = Cm(1.5)
    cat_w     = Cm(round((18.0 - 1.5) / n_cats, 2))  # fits A4 portrait 18cm usable width

    # 3-row table: badge header | male row | female row
    tbl = doc.add_table(rows=3, cols=n_cats + 1)
    tbl.style = "Table Grid"
    _no_table_borders(tbl)

    # ── Row 0: category badge headers ────────────────────────────────────
    r0 = tbl.rows[0]
    _set_cell_bg(r0.cells[0], _NAVY_HEX)
    r0.cells[0].width = label_w
    for j, cat in enumerate(_CATEGORIES):
        c = r0.cells[j + 1]
        c.width = cat_w
        _set_cell_bg(c, _CAT_BADGE_HEX.get(cat, _NAVY_HEX))
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(5)
        run = p.add_run(cat.upper())
        run.bold = True
        run.font.size = Pt(11)
        run.font.color.rgb = _WHITE_RGB

    # ── Rows 1 & 2: male / female champion cells ─────────────────────────
    for row_idx, (label, champs, label_bg) in enumerate(
        [("MALE",   m_champs, _NAVY_HEX),
         ("FEMALE", f_champs, "4a5e6e")],
        start=1,
    ):
        r = tbl.rows[row_idx]

        # Label column
        lc = r.cells[0]
        lc.width = label_w
        _set_cell_bg(lc, label_bg)
        lp = lc.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(8)
        lp.paragraph_format.space_after  = Pt(4)
        lr = lp.add_run(label)
        lr.bold = True
        lr.font.size = Pt(8)
        lr.font.color.rgb = _WHITE_RGB

        # Champion cells
        cell_bg = _WHITE_HEX if row_idx == 1 else _ALT_HEX
        for j, cat in enumerate(_CATEGORIES):
            c = r.cells[j + 1]
            c.width = cat_w
            champ = champs[cat]
            _set_cell_bg(c, cell_bg)

            # Name (bold, navy)
            p1 = c.paragraphs[0]
            p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p1.paragraph_format.space_before = Pt(6)
            p1.paragraph_format.space_after  = Pt(1)
            if champ:
                rn = p1.add_run(champ.name)
                rn.bold = True
                rn.font.size = Pt(10)
                rn.font.color.rgb = _NAVY_RGB

            # Club (small italic, muted)
            p2 = c.add_paragraph()
            p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p2.paragraph_format.space_before = Pt(0)
            p2.paragraph_format.space_after  = Pt(6)
            if champ and champ.preferred_club:
                rc = p2.add_run(champ.preferred_club)
                rc.italic = True
                rc.font.size = Pt(8)
                rc.font.color.rgb = RGBColor(0x70, 0x80, 0x90)


def write_race_report(
    race_num: int,
    total_races: int,
    runners: List[RunnerRaceEntry],
    team_results: List[TeamRaceResult],
    images_dir: Optional[Path],
    year: int,
    filepath: Path,
    source_file: Optional[Path] = None,
) -> None:
    """Write the per-race scoring card DOCX (+ PDF) styled after the WRRL card."""
    race_name = _extract_race_name(source_file) if source_file else f"Race {race_num}"
    filepath  = Path(filepath).with_suffix(".docx")

    doc = Document()
    _apply_document_font_defaults(doc)
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    # Ensure no landscape attribute is present (default template may carry one)
    _pgSz = section._sectPr.find(qn("w:pgSz"))
    if _pgSz is not None:
        _pgSz.attrib.pop(qn("w:orient"), None)

    _build_race_header(doc, race_name, race_num, total_races, images_dir)
    _spacer(doc, height_pt=8)
    _write_race_individual(doc, runners, top_n=10)
    _write_race_category_strip(doc, runners)
    _page_break(doc)
    _write_race_teams(doc, team_results)
    _build_footer(doc, year)

    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))
    log.info("Race %d report written: %s", race_num, filepath.name)

    try:
        from docx2pdf import convert  # type: ignore
        pdf_path = filepath.with_suffix(".pdf")
        convert(str(filepath), str(pdf_path))
        log.info("Race %d PDF written: %s", race_num, pdf_path.name)
    except Exception as exc:
        log.warning("Race %d PDF conversion skipped — %s", race_num, exc)


# ── Document header (branding) ─────────────────────────────────────────────────

def _build_cover_header(
    doc: Document,
    year: int,
    highest_race: int,
    images_dir: Optional[Path],
) -> None:
    """Two-cell table: shield logo left | title text right, both on navy."""
    tbl = doc.add_table(rows=1, cols=2)
    logo_cell  = tbl.rows[0].cells[0]
    title_cell = tbl.rows[0].cells[1]

    logo_cell.width  = Cm(3.0)
    title_cell.width = Cm(15.0)

    _set_cell_bg(logo_cell, _NAVY_HEX)
    _set_cell_bg(title_cell, _NAVY_HEX)

    # Logo
    shield_path = (images_dir / "WRRL shield concept.png") if images_dir else None
    if shield_path and shield_path.exists():
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(shield_path), width=Cm(2.8))

    # Title
    title_cell.paragraphs[0].clear()
    p1 = title_cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p1.paragraph_format.left_indent  = Cm(0.4)
    p1.paragraph_format.space_before = Pt(8)
    r = p1.add_run(f"WRRL League Report  |  Season {year}")
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = _WHITE_RGB

    p2 = title_cell.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.4)
    subtitle = (
        "Wiltshire Road Race League  —  Season Summary"
        if highest_race == 8
        else f"Wiltshire Road Race League  —  Update Race {highest_race}"
    )
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10)
    r2.font.color.rgb = _SUBHDR_RGB

    # Green accent rule below header
    accent = doc.add_paragraph()
    accent.paragraph_format.space_before = Pt(0)
    accent.paragraph_format.space_after  = Pt(6)
    _set_para_bg(accent, _GREEN_HEX)


# ── Footer ─────────────────────────────────────────────────────────────────────

def _build_footer(doc: Document, year: int, include_page_numbers: bool = False) -> None:
    footer = doc.sections[0].footer
    if footer.paragraphs:
        footer.paragraphs[0].clear()

    tbl = footer.add_table(rows=1, cols=3, width=Cm(18.0))
    _no_table_borders(tbl)

    left_cell = tbl.rows[0].cells[0]
    center_cell = tbl.rows[0].cells[1]
    right_cell = tbl.rows[0].cells[2]
    left_cell.width = Cm(6.0)
    center_cell.width = Cm(7.0)
    right_cell.width = Cm(5.0)

    left_p = left_cell.paragraphs[0]
    left_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    left_run = left_p.add_run(f"© {year} Wiltshire Athletics Assoc.")
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = RGBColor(0x80, 0x80, 0x90)

    center_p = center_cell.paragraphs[0]
    center_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    center_run = center_p.add_run("Wiltshire League Scorer v2.1")
    center_run.font.size = Pt(8)
    center_run.font.color.rgb = RGBColor(0x80, 0x80, 0x90)

    if include_page_numbers:
        right_p = right_cell.paragraphs[0]
        right_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        page_label = right_p.add_run("Page ")
        page_label.font.size = Pt(8)
        page_label.font.color.rgb = RGBColor(0x80, 0x80, 0x90)
        _append_page_number(right_p)


# ── Public entry point ─────────────────────────────────────────────────────────

def write_combined_report(
    highest_race: int,
    year: int,
    male_records: List[RunnerSeasonRecord],
    female_records: List[RunnerSeasonRecord],
    div1_teams: List[TeamSeasonRecord],
    div2_teams: List[TeamSeasonRecord],
    unrec_all: List[UnrecognisedClub],
    images_dir: Optional[Path],
    filepath: Path,
) -> None:
    """Write the branded combined report to *filepath* (.docx) and attempt PDF.

    Parameters
    ----------
    highest_race:  highest race number included in this run.
    year:          season year (for title and footer).
    male_records / female_records: season leaderboard records.
    div1_teams / div2_teams: club table records per division.
    unrec_all:     aggregate non-league club list for overview.
    images_dir:    folder containing ``WRRL shield concept.png``.
    filepath:      output path; .docx extension used automatically.
    """
    filepath = Path(filepath).with_suffix(".docx")

    doc = Document()
    _apply_document_font_defaults(doc)

    # ── A4 Portrait ───────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.orientation   = WD_ORIENT.PORTRAIT
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(1.5)
    section.right_margin  = Cm(1.5)
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)

    nonleague_fname = ""  # non-league DOCX no longer generated

    # ── Content ───────────────────────────────────────────────────────────────
    _build_cover_header(doc, year, highest_race, images_dir)
    _write_league_narrative(doc, highest_race, year, male_records, female_records,
                            div1_teams, div2_teams)

    # Club tables
    _spacer(doc, height_pt=4)
    _write_club_table(doc, div1_teams, "Division 1 — Club Table",
                      highest_race, is_div1=True)
    _page_break(doc)
    _write_club_table(doc, div2_teams, "Division 2 — Club Table",
                      highest_race, is_div1=False)

    # Individual scorer section — new page
    _page_break(doc)
    _section_heading(doc, "Individual Runners", space_before_pt=4)
    _write_individual_table(doc, male_records,   "Top 20 — Male",   highest_race, top_n=20)
    doc.add_page_break()
    _write_individual_table(doc, female_records, "Top 20 — Female", highest_race, top_n=20)

    # Category leaders — new page
    doc.add_page_break()
    _write_category_leaders(doc, male_records, female_records, top_n=3)

    # Season overview — final page
    _page_break(doc)
    _write_overview(doc, highest_race, year, male_records, female_records,
                    unrec_all)

    _build_footer(doc, year, include_page_numbers=True)

    # ── Save DOCX ─────────────────────────────────────────────────────────────
    filepath.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(filepath))
    log.info("Report written: %s", filepath.name)

    # ── PDF conversion (requires Microsoft Word on Windows) ───────────────────
    try:
        from docx2pdf import convert  # type: ignore
        pdf_path = filepath.with_suffix(".pdf")
        convert(str(filepath), str(pdf_path))
        log.info("PDF written: %s", pdf_path.name)
    except Exception as exc:
        log.warning("PDF conversion skipped — %s", exc)
