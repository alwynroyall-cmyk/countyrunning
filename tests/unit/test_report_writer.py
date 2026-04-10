from docx import Document
from docx.oxml.ns import qn

from league_scorer.report_writer import (
    _cell_text,
    _section_heading,
    _set_cell_bg,
    _set_para_bg,
)


def test_set_cell_bg_applies_fill_to_table_cell():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    _set_cell_bg(cell, "abcdef")

    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "abcdef"


def test_set_para_bg_applies_fill_to_paragraph():
    doc = Document()
    para = doc.add_paragraph("Test")

    _set_para_bg(para, "123456")

    pPr = para._p.get_or_add_pPr()
    shd = pPr.find(qn("w:shd"))
    assert shd is not None
    assert shd.get(qn("w:fill")) == "123456"


def test_cell_text_sets_run_text_and_style():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)

    _cell_text(cell, "Hello", bold=True, italic=True, size_pt=12)

    run = cell.paragraphs[0].runs[0]
    assert run.text == "Hello"
    assert run.bold is True
    assert run.italic is True
    assert run.font.size.pt == 12


def test_section_heading_adds_heading_paragraph():
    doc = Document()
    _section_heading(doc, "Season Overview", space_before_pt=5)

    para = doc.paragraphs[0]
    assert para.text == "SEASON OVERVIEW"
    assert para.paragraph_format.space_before.pt == 5
